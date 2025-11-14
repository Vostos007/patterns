"""DOCX export helpers for KPS pipeline."""

from __future__ import annotations

import logging
from difflib import SequenceMatcher
from html import unescape
from pathlib import Path
from typing import Iterable, List, Optional

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from kps.core.document import BlockType, KPSDocument, SectionType

logger = logging.getLogger(__name__)


def render_docx_inplace(
    original_docx: Path,
    original_doc: KPSDocument,
    translated_doc: KPSDocument,
    output_path: Path,
) -> Path:
    """Render translated DOCX by replacing text inside the original template."""

    document = Document(str(original_docx))
    paragraphs = list(_iter_paragraphs(document))

    original_texts = _collect_block_text(original_doc)
    translated_texts = _collect_block_text(translated_doc)

    if len(original_texts) != len(translated_texts):
        logger.warning(
            "Block count mismatch: original=%s translated=%s",
            len(original_texts),
            len(translated_texts),
        )

    paragraph_groups = _group_blocks_by_paragraph(paragraphs, original_texts)

    # SAFETY: ensure the very first block (usually document title) maps to the first paragraph
    if paragraphs and original_texts:
        first_block_idx = 0
        current_owner = None
        for idx, group in enumerate(paragraph_groups):
            if first_block_idx in group:
                current_owner = idx
                break
        if current_owner is None:
            paragraph_groups[0].insert(0, first_block_idx)
        elif current_owner != 0:
            paragraph_groups[current_owner].remove(first_block_idx)
            paragraph_groups[0].insert(0, first_block_idx)

    if logger.isEnabledFor(logging.DEBUG):
        for idx in range(min(3, len(paragraphs))):
            logger.debug(
                "Paragraph[%s] text=%r blocks=%s",
                idx,
                paragraphs[idx].text if idx < len(paragraphs) else "",
                paragraph_groups[idx] if idx < len(paragraph_groups) else [],
            )

    updated = 0
    skipped = 0
    for paragraph, block_indexes in zip(paragraphs, paragraph_groups):
        if not block_indexes:
            skipped += 1
            continue
        new_text = "\n".join(translated_texts[i] for i in block_indexes if translated_texts[i])
        if not new_text:
            skipped += 1
            continue
        _replace_paragraph_text(paragraph, new_text)
        updated += 1

    if skipped:
        logger.warning("Skipped %s blocks because matching paragraph not found", skipped)
    logger.info("Updated %s paragraphs in DOCX", updated)

    _rewrite_tables(document, translated_doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path


def _collect_block_text(doc: KPSDocument) -> List[str]:
    texts: List[str] = []
    for section in doc.sections:
        for block in section.blocks:
            texts.append(block.content or "")
    return texts


def _iter_block_items(parent) -> Iterable[object]:
    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:  # Table, etc.
        parent_elm = parent._tbl
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _iter_paragraphs(parent) -> Iterable[Paragraph]:
    for block in _iter_block_items(parent):
        if isinstance(block, Paragraph):
            yield block


def _align_paragraphs(
    paragraphs: List[Paragraph], original_texts: List[str]
) -> List[Optional[Paragraph]]:
    matches: List[Optional[Paragraph]] = []
    idx = 0
    total = len(paragraphs)
    for original_text in original_texts:
        norm = _normalize(original_text)
        if not norm:
            matches.append(None)
            continue
        found = None
        search_idx = idx
        while search_idx < total:
            para = paragraphs[search_idx]
            para_norm = _normalize(para.text)
            if _paragraph_matches(para_norm, norm):
                found = para
                idx = search_idx + 1
                break
            search_idx += 1
        matches.append(found)
    return matches


def _group_blocks_by_paragraph(
    paragraphs: List[Paragraph], original_texts: List[str]
) -> List[List[int]]:
    """
    Map sequential Docling blocks onto DOCX paragraphs, allowing multiple blocks per
    paragraph when Docling splits sentences across nodes.
    """

    total_blocks = len(original_texts)
    total_paragraphs = len(paragraphs)
    groups: List[List[int]] = [[] for _ in range(total_paragraphs)]
    normalized_paragraphs = [_normalize(p.text) for p in paragraphs]
    accumulators = ["" for _ in range(total_paragraphs)]

    block_idx = 0
    para_idx = 0

    while block_idx < total_blocks and para_idx < total_paragraphs:
        para_norm = normalized_paragraphs[para_idx]
        if not para_norm:
            para_idx += 1
            continue

        block_text = original_texts[block_idx]
        block_idx += 1
        block_norm = _normalize(block_text)
        if not block_norm:
            continue

        groups[para_idx].append(block_idx - 1)
        accumulators[para_idx] = (accumulators[para_idx] + " " + block_norm).strip()
        accum_norm = accumulators[para_idx]

        if _paragraph_matches(accum_norm, para_norm) or _paragraph_matches(
            para_norm, accum_norm
        ):
            para_idx += 1
        elif len(accum_norm) >= len(para_norm) * 1.5:
            # Prevent runaway accumulation when Docling order drifts
            logger.debug(
                "Paragraph %s drifted (accum len=%s, para len=%s)",
                para_idx,
                len(accum_norm),
                len(para_norm),
            )
            para_idx += 1

    if block_idx < total_blocks:
        logger.warning(
            "Unmapped doc blocks remaining after DOCX alignment: %s",
            total_blocks - block_idx,
        )

    return groups


def _replace_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if not paragraph.runs:
        paragraph.add_run("")
    paragraph.runs[0].text = new_text


def _normalize(text: str) -> str:
    if not text:
        return ""
    text = unescape(text)
    text = text.replace("\xa0", " ")
    text = text.replace("_", " ")
    text = text.lower()
    return " ".join(text.split())


def _paragraph_matches(para_norm: str, block_norm: str) -> bool:
    """Fuzzy paragraph match that tolerates Docling block reshaping."""

    if para_norm == block_norm:
        return True

    if not para_norm or not block_norm:
        return False

    if block_norm in para_norm:
        return True

    if para_norm in block_norm:
        return True

    # Fall back to similarity threshold for minor punctuation/spacing drift.
    ratio = SequenceMatcher(None, para_norm, block_norm).ratio()
    return ratio >= 0.92


def build_docx_from_structure(translated_doc: KPSDocument, output_path: Path) -> Path:
    """Create a fresh DOCX from translated blocks when original template is unavailable."""

    logger.info("Building DOCX from structure → %s", output_path)
    doc = Document()
    if doc.paragraphs:
        _delete_paragraph(doc.paragraphs[0])
    else:
        doc.add_paragraph("")

    for section in translated_doc.sections:
        first_block = section.blocks[0] if section.blocks else None
        heading_inserted = False

        if section.title:
            heading_text = section.title
            use_first_block = (
                first_block
                and first_block.content
                and (
                    first_block.block_type == BlockType.HEADING
                    or section.section_type == SectionType.COVER
                )
            )
            if use_first_block:
                heading_text = first_block.content
                heading_inserted = True
            doc.add_heading(heading_text, level=2)
        elif first_block and first_block.block_type == BlockType.HEADING:
            doc.add_heading(first_block.content or "", level=3)
            heading_inserted = True

        for idx, block in enumerate(section.blocks):
            if heading_inserted and idx == 0:
                continue
            _append_block(doc, block.block_type, block.content)

    # Remove default placeholder title paragraphs that Word templates add
    for para in list(doc.paragraphs):
        try:
            style_name = para.style.name if para.style else ""
        except AttributeError:
            style_name = ""
        text_value = para.text.strip()
        if style_name == "Title" or text_value.endswith("_docx"):
            logger.debug(
                "Removing placeholder paragraph: %r (style=%s)",
                text_value,
                style_name,
            )
            _delete_paragraph(para)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    _cleanup_docx_placeholders(output_path)
    return output_path


def _append_block(doc: Document, block_type: BlockType, content: str) -> None:
    text = content or ""
    if block_type == BlockType.HEADING:
        doc.add_heading(text, level=3)
    elif block_type == BlockType.LIST:
        for line in text.splitlines():
            if line.strip():
                doc.add_paragraph(line.strip(), style="List Bullet")
    elif block_type == BlockType.TABLE:
        rows = [line.split("\t") for line in text.splitlines() if line]
        if not rows:
            doc.add_paragraph(text)
            return
        cols = max(len(row) for row in rows)
        table = doc.add_table(rows=len(rows), cols=cols)
        for r_idx, row in enumerate(rows):
            for c_idx, value in enumerate(row):
                cell = table.cell(r_idx, c_idx)
                _clear_paragraph(cell.paragraphs[0])
                cell.paragraphs[0].add_run(value)
    elif block_type == BlockType.FIGURE:
        para = doc.add_paragraph()
        para.style = "Caption"
        para.add_run(text)
    else:
        for idx, line in enumerate(text.splitlines() or [""]):
            if idx == 0:
                para = doc.add_paragraph()
            else:
                para = doc.add_paragraph()
            para.add_run(line)


def _clear_paragraph(paragraph: Paragraph) -> None:
    for run in paragraph.runs:
        run.text = ""


def _delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _cleanup_docx_placeholders(docx_path: Path) -> None:
    """Remove leftover Title paragraphs after saving."""
    try:
        doc = Document(str(docx_path))
    except Exception as exc:
        logger.debug("Unable to reopen DOCX for cleanup: %s", exc)
        return

    removed = False
    for para in list(doc.paragraphs):
        style_name = para.style.name if para.style else ""
        text_value = para.text.strip()
        if style_name == "Title" or text_value.endswith("_docx"):
            _delete_paragraph(para)
            removed = True

    if removed:
        doc.save(str(docx_path))
    logger.info("DOCX placeholder cleanup: removed=%s path=%s", removed, docx_path)


def _rewrite_tables(document: Document, translated_doc: KPSDocument) -> None:
    """Rewrite DOCX tables using translated table blocks."""

    table_blocks: List[str] = []
    for section in translated_doc.sections:
        for block in section.blocks:
            if block.block_type == BlockType.TABLE:
                table_blocks.append(block.content or "")

    if not table_blocks:
        return

    block_iter = iter(table_blocks)
    replaced = 0
    for table in document.tables:
        try:
            block_text = next(block_iter)
        except StopIteration:
            block_text = None

        if not block_text:
            continue

        rows = _parse_table_block(block_text)
        if not rows:
            continue

        if not _table_contains_cyrillic(table):
            continue

        max_cols = max(len(row) for row in rows)
        for r_idx, row in enumerate(rows):
            if r_idx >= len(table.rows):
                break
            for c_idx in range(max_cols):
                if c_idx >= len(table.rows[r_idx].cells):
                    break
                value = row[c_idx] if c_idx < len(row) else ""
                cell = table.rows[r_idx].cells[c_idx]
                paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph("")
                _clear_paragraph(paragraph)
                paragraph.add_run(value)
        replaced += 1

    remaining = sum(1 for _ in block_iter)
    if remaining:
        logger.warning(
            "%s translated tables were not applied because matching DOCX tables were not found",
            remaining,
        )
    if replaced:
        logger.info("Updated %s DOCX tables via template pipeline", replaced)


def _table_contains_cyrillic(table: Table) -> bool:
    import re

    pattern = re.compile(r"[А-Яа-яЁё]")
    for row in table.rows:
        for cell in row.cells:
            if pattern.search(cell.text):
                return True
    return False


def _parse_table_block(block_text: str) -> List[List[str]]:
    rows: List[List[str]] = []
    for raw_line in block_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        stripped_line = line.strip("| ")
        if not stripped_line:
            continue
        if all(ch in "-|: " for ch in stripped_line):
            continue

        if "\t" in raw_line:
            cells = [cell.strip() for cell in raw_line.split("\t")]
        else:
            cells = [cell.strip() for cell in stripped_line.split("|")]

        rows.append(cells)

    return rows
