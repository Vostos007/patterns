"""
Create reference.docx template with proper styles for Pandoc.

This script creates a minimal reference.docx file with:
- Heading styles (1-6)
- Normal paragraph style
- Code style
- Table style
- Caption style
- Custom fonts and spacing
"""

from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx не установлен. Установите: pip install python-docx")
    exit(1)


def create_reference_docx(output_path: str):
    """Create reference.docx with custom styles."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.15

    # Heading 1 style
    h1_style = doc.styles["Heading 1"]
    h1_font = h1_style.font
    h1_font.name = "Calibri Light"
    h1_font.size = Pt(16)
    h1_font.color.rgb = RGBColor(46, 116, 181)
    h1_font.bold = False
    h1_style.paragraph_format.space_before = Pt(12)
    h1_style.paragraph_format.space_after = Pt(6)

    # Heading 2 style
    h2_style = doc.styles["Heading 2"]
    h2_font = h2_style.font
    h2_font.name = "Calibri Light"
    h2_font.size = Pt(13)
    h2_font.color.rgb = RGBColor(46, 116, 181)
    h2_font.bold = False
    h2_style.paragraph_format.space_before = Pt(10)
    h2_style.paragraph_format.space_after = Pt(4)

    # Heading 3 style
    h3_style = doc.styles["Heading 3"]
    h3_font = h3_style.font
    h3_font.name = "Calibri Light"
    h3_font.size = Pt(12)
    h3_font.color.rgb = RGBColor(46, 116, 181)
    h3_font.bold = False

    # Code style (create if doesn't exist)
    try:
        code_style = doc.styles["Code"]
    except KeyError:
        code_style = doc.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)

    code_font = code_style.font
    code_font.name = "Consolas"
    code_font.size = Pt(9)
    code_font.color.rgb = RGBColor(0, 0, 0)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)

    # Caption style
    try:
        caption_style = doc.styles["Caption"]
    except KeyError:
        caption_style = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)

    caption_font = caption_style.font
    caption_font.name = "Calibri"
    caption_font.size = Pt(9)
    caption_font.italic = True
    caption_font.color.rgb = RGBColor(68, 68, 68)
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add sample content to demonstrate styles
    doc.add_heading("Reference Document Template", 0)

    p = doc.add_paragraph(
        "This is a reference document for Pandoc DOCX export. "
        "All styles defined here will be used when rendering documents."
    )

    doc.add_heading("Heading 1 Example", 1)
    doc.add_paragraph("This is normal paragraph text under Heading 1.")

    doc.add_heading("Heading 2 Example", 2)
    doc.add_paragraph("This is normal paragraph text under Heading 2.")

    doc.add_heading("Heading 3 Example", 3)
    doc.add_paragraph("This is normal paragraph text under Heading 3.")

    # Add code block
    doc.add_paragraph("Code block example:", style="Normal")
    code_para = doc.add_paragraph(
        "def hello_world():\n    print('Hello, World!')", style="Code"
    )

    # Add table
    doc.add_paragraph("Table example:", style="Normal")
    table = doc.add_table(rows=3, cols=3)
    table.style = "Light Grid Accent 1"

    # Fill table
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Column 1"
    hdr_cells[1].text = "Column 2"
    hdr_cells[2].text = "Column 3"

    for i in range(1, 3):
        row_cells = table.rows[i].cells
        row_cells[0].text = f"Row {i}, Col 1"
        row_cells[1].text = f"Row {i}, Col 2"
        row_cells[2].text = f"Row {i}, Col 3"

    # Add caption
    doc.add_paragraph("Table 1: Example table with caption", style="Caption")

    # Save
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"✓ Created reference.docx: {output}")


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python create_reference_docx.py OUTPUT_PATH")
        print("Example: python create_reference_docx.py configs/reference.docx")
        sys.exit(1)

    create_reference_docx(sys.argv[1])
