"""
End-to-end test for DOCX export with styles.

Verifies that:
- DOCX file is created
- File contains actual content
- Styles from reference.docx are applied
"""

from pathlib import Path
import pytest


def test_docx_file_exists():
    """Test that DOCX output file exists and is not empty."""
    docx_path = Path("output/final.docx")

    assert docx_path.exists(), f"DOCX file not found: {docx_path}"
    assert docx_path.stat().st_size > 0, "DOCX file is empty"


def test_docx_has_styles():
    """Test that DOCX contains expected styles from reference.docx."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    docx_path = Path("output/final.docx")

    if not docx_path.exists():
        pytest.skip(f"DOCX file not found: {docx_path}")

    doc = Document(str(docx_path))

    # Check that document has paragraphs
    assert len(doc.paragraphs) > 0, "DOCX has no paragraphs"

    # Check for heading styles
    has_heading = any(
        p.style and p.style.name and p.style.name.startswith("Heading")
        for p in doc.paragraphs
    )
    assert has_heading, "No Heading styles found in DOCX"

    # Check for normal text
    has_text = any(p.style and p.style.name == "Normal" for p in doc.paragraphs)
    assert has_text, "No Normal style paragraphs found in DOCX"


def test_docx_has_content():
    """Test that DOCX contains actual text content."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    docx_path = Path("output/final.docx")

    if not docx_path.exists():
        pytest.skip(f"DOCX file not found: {docx_path}")

    doc = Document(str(docx_path))

    # Get all text
    all_text = "\n".join([p.text for p in doc.paragraphs])

    assert len(all_text.strip()) > 0, "DOCX contains no text"
    assert len(all_text) > 100, "DOCX text content is suspiciously short"


def test_docx_table_style():
    """Test that DOCX tables have proper styling."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    docx_path = Path("output/final.docx")

    if not docx_path.exists():
        pytest.skip(f"DOCX file not found: {docx_path}")

    doc = Document(str(docx_path))

    # If document has tables, check their styles
    if doc.tables:
        for table in doc.tables:
            # Table should have a style
            assert table.style is not None, "Table has no style"
            # Table should have rows and cells
            assert len(table.rows) > 0, "Table has no rows"
            if table.rows:
                assert len(table.rows[0].cells) > 0, "Table has no cells"
