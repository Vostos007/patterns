import shutil

import pytest

from docx import Document

from kps.export.docling_renderer import render_docx, render_markdown
from kps.export.docling_writer import apply_translations
from kps.extraction.docling_extractor import DoclingExtractor
from kps.extraction.segmenter import Segmenter


def _make_sample_docx(tmp_path):
    sample = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Пример текста")
    doc.add_paragraph("Вторая строка")
    doc.save(sample)
    return sample


def _prepare_docling(tmp_path):
    path = _make_sample_docx(tmp_path)
    extractor = DoclingExtractor()
    kps_doc = extractor.extract_document(path, slug="sample")
    segments = Segmenter().segment_document(kps_doc)
    assert kps_doc.docling_document is not None
    return kps_doc.docling_document, segments


def test_docling_markdown_export(tmp_path):
    docling_doc, segments = _prepare_docling(tmp_path)
    translations = [f"Line {idx}" for idx, _ in enumerate(segments)]
    updated, missing = apply_translations(docling_doc, segments, translations)
    assert not missing

    md_path = tmp_path / "out.md"
    render_markdown(updated, md_path)
    content = md_path.read_text(encoding="utf-8")
    assert "Line 0" in content
    assert "Пример" not in content


@pytest.mark.skipif(shutil.which("pandoc") is None, reason="pandoc required")
def test_docling_docx_export(tmp_path):
    docling_doc, segments = _prepare_docling(tmp_path)
    translations = ["Hello" for _ in segments]
    updated, missing = apply_translations(docling_doc, segments, translations)
    assert not missing

    out_path = tmp_path / "out.docx"
    render_docx(updated, out_path)
    doc = Document(out_path)
    texts = [p.text for p in doc.paragraphs]
    assert any("Hello" in t for t in texts)
    assert all("Пример" not in t for t in texts)
